#---------------------------------------------------------------------------------------------------------------------------------
### Authenticator
#---------------------------------------------------------------------------------------------------------------------------------
import streamlit as st
#---------------------------------------------------------------------------------------------------------------------------------
### Import Libraries
#---------------------------------------------------------------------------------------------------------------------------------

#----------------------------------------
import numpy as np
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
#----------------------------------------
import os
import sys
import io
import traceback
from PIL import Image
#----------------------------------------
from io import BytesIO
#----------------------------------------
import fitz 
import pikepdf
import pdfplumber
from docx import Document
#from pdf2docx import Converter
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
from pdf2image import convert_from_path, convert_from_bytes
from pdf2image.exceptions import PDFInfoNotInstalledError,PDFPageCountError,PDFSyntaxError
#from pdfminer.high_level import extract_text
#----------------------------------------
#import nltk
#nltk.download('punkt')
#nltk.download('punkt_tab')
#---------------------------------------------------------------------------------------------------------------------------------
### Title and description for your Streamlit app
#---------------------------------------------------------------------------------------------------------------------------------
st.set_page_config(page_title="PDF Playground | v0.2",
                    layout="wide",
                    page_icon="📘",            
                    initial_sidebar_state="collapsed")
#---------------------------------------------------------------------------------------------------------------------------------
### Login Page | Streamlit app
#---------------------------------------------------------------------------------------------------------------------------------

             
#---------------------------------------------------------------------------------------------------------------------------------
### CSS
#---------------------------------------------------------------------------------------------------------------------------------
st.markdown("""
    <style>
    .centered-info {
        display: flex;
        justify-content: center;
        align-items: center;
        font-weight: bold;
        font-size: 15px;
        color: #007BFF; 
        background-color: #FFFFFF; 
        border-radius: 5px;
        border: 1px solid #007BFF;
        margin: 0px;
        padding: 5px 10px;
    }
    </style>
""", unsafe_allow_html=True)

st.markdown(
    """
    <style>
    /* Card styling (matches your screenshot) */
    .card {
        background-color: white;
        border-radius: 10px;
        padding: 10px;
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.05);
        transition: all 0.3s ease;
        border: 1px solid #e1e5eb;
        height: 100%;
        display: flex;
        flex-direction: column;
    }
    .card:hover {
        transform: translateY(-4px);
        box-shadow: 0 8px 20px rgba(0, 0, 0, 0.1);
        border-color: #66b2ff;
    }
    .card-title {
        font-size: 1.4rem;
        font-weight: 600;
        color: #1a365d;
        margin-bottom: 16px;
        display: flex;
        align-items: center;
        gap: 8px;
    }
    .card-icon {
        font-size: 1.8rem;
    }
    .card-list {
        padding-left: 0;
        margin-top: 12px;
    }
    .card-list li {
        margin-bottom: 8px;
        font-size: 0.95rem;
        color: #4a5568;
    }
    .card-list li::before {
        content: "✅";
        margin-right: 8px;
        color: #3182ce;
    }

    /* Home button style (top-left) */
    .home-btn {
        background: #f7fafc;
        border: 1px solid #e2e8f0;
        border-radius: 8px;
        padding: 6px 16px;
        font-size: 14px;
        font-weight: 500;
        color: #2d3748;
        display: inline-flex;
        align-items: center;
        gap: 6px;
        cursor: pointer;
        transition: all 0.2s;
    }
    .home-btn:hover {
        background: #edf2f7;
        color: #1a365d;
    }

    /* Sidebar-like input panel */
    .input-panel {
        background: #f8fafc;
        border-radius: 10px;
        padding: 20px;
        border: 1px solid #e2e8f0;
    }
    </style>
    """,
    unsafe_allow_html=True
)

#---------------------------------------------------------------------------------------------------------------------------------
### Description for your Streamlit app
#---------------------------------------------------------------------------------------------------------------------------------
st.markdown(
    """
    <style>
    .title-large {
        text-align: center;
        font-size: 35px;
        font-weight: bold;
        background: linear-gradient(to left, red, orange, blue, indigo, violet);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .title-small {
        text-align: center;
        font-size: 20px;
        background: linear-gradient(to left, red, orange, blue, indigo, violet);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .version-badge {
        text-align: center;
        display: inline-block;
        background: linear-gradient(120deg, #0056b3, #0d4a96);
        color: white;
        padding: 2px 12px;
        border-radius: 20px;
        font-size: 1.15rem;
        margin-top: 8px;
        font-weight: 600;
        letter-spacing: 0.5px;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.15);
    }
    </style>
    <div style="text-align: center;">
        <div class="title-large">PDF Playground</div>
        <div class="version-badge"> Play with PDF | v0.2 </div>
    </div>
    """,
    unsafe_allow_html=True
)

#----------------------------------------
st.markdown(
    """
    <style>
    .footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background-color: #F0F2F6;
        text-align: center;
        padding: 10px;
        font-size: 14px;
        color: #333;
        z-index: 100;
    }
    .footer p {
        margin: 0;
    }
    .footer .highlight {
        font-weight: bold;
        color: blue;
    }
    </style>

    <div class="footer">
        <p>© 2026 | Created by : <span class="highlight">Avijit Chakraborty</span> <a href="mailto:avijit.mba18@gmail.com"> 📩 </a> | <span class="highlight">Thank you for visiting the app | Unauthorized uses or copying is strictly prohibited | For best view of the app, please zoom out the browser to 75%.</span> </p>
    </div>
    """,
    unsafe_allow_html=True)

#---------------------------------------------------------------------------------------------------------------------------------
### Functions & Definitions
#---------------------------------------------------------------------------------------------------------------------------------

def pdf_to_images(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    images = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(img)
    return images

def merge_pdfs(pdf_files):
    merger = PdfMerger()
    for pdf_file in pdf_files:
        merger.append(pdf_file)
    merged_pdf = io.BytesIO()
    merger.write(merged_pdf)
    merger.close()
    merged_pdf.seek(0)
    return merged_pdf

def compress_pdf(input_pdf, compression_factor=0.5):
    """Compress a PDF file and return it as bytes."""
    
    doc = fitz.open(stream=input_pdf.read(), filetype="pdf")
    output_stream = io.BytesIO()
    new_doc = fitz.open()  # Create a new PDF

    for page in doc:
        # Convert PDF page to Pixmap
        pix = page.get_pixmap()  
        
        # Convert Pixmap to PIL Image
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        
        # Compress the image and store in a BytesIO stream
        img_stream = io.BytesIO()
        img.save(img_stream, format="JPEG", quality=int(compression_factor * 100))
        img_stream.seek(0)

        # Proper way to create a new Pixmap from a compressed JPEG
        compressed_pix = fitz.Pixmap(fitz.csRGB, fitz.open("jpeg", img_stream).extract_image(0)["image"])

        # Create new page and insert compressed image
        new_page = new_doc.new_page(width=page.rect.width, height=page.rect.height)
        new_page.insert_image(page.rect, pixmap=compressed_pix)

    # Save the new compressed PDF
    new_doc.save(output_stream)
    new_doc.close()

    output_stream.seek(0)
    return output_stream.getvalue()  # Return compressed PDF as bytes

def pdf_to_images_bytes(pdf_bytes):
    images = convert_from_bytes(pdf_bytes)
    return images

def extract_metadata(pdf_file):
    pdf_reader = PdfReader(pdf_file)
    metadata = pdf_reader.metadata
    return metadata

def extract_text(pdf_file):
    pdf_reader = PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def convert_pdf_to_images(pdf_bytes):
    return convert_from_bytes(pdf_bytes)

def extract_text_from_pdf(uploaded_file, start_page=None, end_page=None):
    """Extract text from an entire PDF, a specific page, or a range of pages."""
    text = ""
    with pdfplumber.open(uploaded_file) as pdf:
        total_pages = len(pdf.pages)
        start_page = max(0, (start_page - 1) if start_page else 0)
        end_page = min(total_pages, end_page if end_page else total_pages)

        for i in range(start_page, end_page):
            extracted_text = pdf.pages[i].extract_text()
            if extracted_text:
                text += f"\n--- Page {i+1} ---\n"  # Properly format text with page numbers
                text += extracted_text + "\n"
    return text.strip()

def summarize_text(text, num_sentences):
    """Summarizes extracted text using LSA algorithm."""
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = LsaSummarizer()
    summary_sentences = summarizer(parser.document, num_sentences)
    summary = "\n".join(str(sentence) for sentence in summary_sentences)
    return summary

#---------------------------------------------------------------------------------------------------------------------------------
### Main app
#---------------------------------------------------------------------------------------------------------------------------------

#---------------------------------------------------------------------------------------------------------------------------------
#---------------------------------------------------------------------------------------------------------------------------------
                      
# Initialize session state
if 'page' not in st.session_state:
    st.session_state.page = 'home'

def go_home():
    st.session_state.page = 'home'

#---------------------------------------------------------------------------------------------------------------------------------
if st.session_state.page == 'home':

    st.markdown("""
    <style>
    .banner {
        background: linear-gradient(135deg, #f0f7ff 0%, #e6f2ff 100%);
        border-radius: 32px;
        padding: 15px;
        margin: 25px 0;
        border: 1px solid rgba(0, 86, 179, 0.15);
        text-align: center;
        font-size: 1.15rem;
        color: #0056b3;
        font-weight: 600;
    }
    </style>

    <div class="banner">
        Click the cards below to access different sections and explore the following features
    </div>
    """, unsafe_allow_html=True)

    # Helper function to create a uniform card to avoid code repetition
    def render_tool_card(col, title, icon, description, page_key, button_key):
        with col:
            st.markdown(f"""
            <div class="card" style="
                border: 1px solid #e0e0e0; 
                border-radius: 12px; 
                padding: 20px; 
                height: 100%; 
                background-color: #ffffff;
                box-shadow: 0 2px 5px rgba(0,0,0,0.05);
                transition: transform 0.2s, box-shadow 0.2s;
                display: flex;
                flex-direction: column;
                justify-content: space-between;
            ">
                <div>
                    <div class="card-title" style="
                        font-size: 1.2em; 
                        font-weight: 700; 
                        margin-bottom: 12px; 
                        color: #2c3e50;
                        display: flex;
                        align-items: center;
                        gap: 8px;
                    ">
                        <span style="font-size: 1.4em;">{icon}</span> {title}
                    </div>
                    <ul class="card-list" style="
                        list-style: none; 
                        padding: 0; 
                        margin: 0; 
                        font-size: 0.9em; 
                        color: #555;
                        line-height: 1.5;
                    ">
                        <li>{description}</li>
                    </ul>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            if st.button("**🚀 Click to Enter**", key=button_key, use_container_width=True, type="primary"):
                st.session_state.page = page_key
                st.rerun()

    # --- Row 1: First 5 Tools ---
    cols_row1 = st.columns(5)

    render_tool_card(cols_row1[0], "View", "👁️", "Preview PDF files directly within the application.", 'pdf_view', 'btn_pdf_view')
    render_tool_card(cols_row1[1], "Extract", "📤", "Extract text and metadata from PDF files efficiently.", 'pdf_ext', 'btn_pdf_ext')
    render_tool_card(cols_row1[2], "Merge", "🔗", "Combine multiple PDF files into a single document.", 'pdf_mer', 'btn_pdf_mer')
    render_tool_card(cols_row1[3], "Compress", "🗜️", "Reduce the file size of uploaded PDF documents.", 'pdf_comp', 'btn_pdf_comp')
    render_tool_card(cols_row1[4], "Protect", "🔒", "Add password protection to secure your PDF files.", 'pdf_pro', 'btn_pdf_pro')

    st.write("") # Spacer between rows

    # --- Row 2: Remaining 4 Tools ---
    # We use 5 columns again but leave the last one empty to center the 4 cards, 
    # OR just use 4 columns. Let's use 4 columns for better width usage.
    cols_row2 = st.columns(4)

    render_tool_card(cols_row2[0], "Unlock", "🔓", "Remove password protection from PDF files.", 'pdf_un', 'btn_pdf_un')
    render_tool_card(cols_row2[1], "Rotate", "🔄", "Change the orientation of pages within a PDF file.", 'pdf_rot', 'btn_pdf_rot')
    render_tool_card(cols_row2[2], "Convert", "📐", "Conversion of a PDF file dynamically.", 'pdf_con', 'btn_pdf_con')
    render_tool_card(cols_row2[3], "Summarize", "📝", "Generate summaries of uploaded PDFs without Generative AI.", 'pdf_sum', 'btn_pdf_sum')
                                                    
#---------------------------------------------------------------------------------------------------------------------------------
### View
#---------------------------------------------------------------------------------------------------------------------------------

#with tab1:
#if page == "view":
elif st.session_state.page == 'pdf_view':
            
            st.markdown("---")      
            col_home, title= st.columns([2,15,])
            
            with col_home:
                if st.button("Home", icon="🏠", key="home_fd", type="secondary", use_container_width=True):
                    go_home()
                    st.rerun()

            with title:
                st.info("""The **View** tab allows you to preview PDF files directly within the application. You can upload a PDF and view its content without any external software.""")
            
            col1, col2, col3 = st.columns((0.2,0.6,0.2))
            with col1:           
                with st.container(border=True):
                    uploaded_file = st.file_uploader("**:blue[Choose PDF file]**",type="pdf",key="file_uploader_preview")
                    if uploaded_file is not None:
                        st.success("PDFs loaded successfully!")

                        try:
                            with col2:
                                with st.container(border=True, height=800):
                                    
                                    images = pdf_to_images(uploaded_file)
                                    if images and isinstance(images, list):
                                        for i, image in enumerate(images):
                                            st.image(image, caption=f'Page {i + 1}', use_container_width=True)
                                    else:
                                        st.warning("No images were generated from the PDF.")

                            with col3:
                                with st.container(border=True):

                                    metadata = extract_metadata(uploaded_file)
                                    if metadata:
                                        metadata_df = pd.DataFrame(list(metadata.items()), columns=["Key", "Value"])
                                        st.table(metadata_df)
                                    else:
                                        st.write("No metadata found in the PDF file.")
                        except Exception as e:
                            st.error(f"An error occurred while processing the PDF: {e}")
                    else:
                        st.warning("Please upload a PDF file to view.")

#---------------------------------------------------------------------------------------------------------------------------------
### Extract
#---------------------------------------------------------------------------------------------------------------------------------

#with tab2:
#if page == "extract":
elif st.session_state.page == 'pdf_ext':

            st.markdown("---")
            col_home, title= st.columns([2,15,])
            
            with col_home:
                if st.button("Home", icon="🏠", key="home_fd", type="secondary", use_container_width=True):
                    go_home()
                    st.rerun()

            with title:
                st.info("""The **Extract** tab is designed to extract text and metadata from PDF files. You can upload a PDF, and the tool will pull out the text content for further analysis or editing.""")
       
            col1, col2 = st.columns((0.2,0.8))
            with col1:           
                with st.container(border=True):            
                    uploaded_file = st.file_uploader("**:blue[Choose PDF file]**", type="pdf", key="file_uploader_extract")
                    if uploaded_file is not None:
                        st.success("PDFs loaded successfully!")
                        with pdfplumber.open(uploaded_file) as pdf:
                            total_pages = len(pdf.pages)
                            
                        #st.success("Text extracted successfully!")
                        #st.write(f"You have selected **{uploaded_file.name}** for text extraction.")
                        st.write(f"📄 **Total Pages in PDF:** {total_pages}")
                        st.divider()
                        
                        extract_type = st.radio("**:blue[Extraction Type]**", ["Text", "Tables"],horizontal=True)
                        page_selection = st.radio("**:blue[Extract from]**", ["All Pages", "Specific Page"],horizontal=True)
                        if page_selection == "Specific Page":
                            selected_page = st.number_input("**:blue[Enter page number]**", min_value=1, max_value=total_pages, value=1)

                        if st.button("**Extract**"):
                            with col2:
                                with st.container(border=True):
                                    
                                    try:
                                        with pdfplumber.open(uploaded_file) as pdf:
                                            extracted_data = ""
                        
                                            if extract_type == "Text":
                                                with st.spinner("Extracting text..."):
                                                    if page_selection == "All Pages":
                                                        extracted_data = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
                                                    else:
                                                        extracted_data = pdf.pages[selected_page - 1].extract_text()

                                                st.text_area("Extracted Text", value=extracted_data, height=600)
                                                text_file = io.BytesIO(extracted_data.encode("utf-8"))


                                            elif extract_type == "Tables":
                                                with st.spinner("Extracting tables..."):
                                                    tables = []
                                                    if page_selection == "All Pages":
                                                        for page in pdf.pages:
                                                            tables.extend(page.extract_tables())
                                                    else:
                                                        tables = pdf.pages[selected_page - 1].extract_tables()

                                                    if tables:
                                                        for i, table in enumerate(tables):
                                                            st.write(f"**Table {i+1}:**")
                                                            st.table(table)
                                            
                                                        csv_data = "\n".join([",".join(map(str, row)) for table in tables for row in table])
                                                        csv_file = io.BytesIO(csv_data.encode("utf-8"))
                                                        
                                             
                                                    else:
                                                        st.warning("⚠️ No tables found in the selected page(s).")                      
                                                    
                                    except Exception as e:
                                        st.error(f"An error occurred while processing the PDF: {e}")
                                
                                if extract_type == "Text":    
                                    st.download_button(label="**📥 Download Extracted Text (.txt)**",data=text_file,file_name="extracted_text.txt",mime="text/plain")   
                                elif extract_type == "Tables":
                                    st.download_button(label="**📥 Download Extracted Tables (.csv)**",data=csv_file,file_name="extracted_tables.csv",mime="text/csv")
                                        
                    else:
                        st.warning("Please upload a PDF file to extract.")

#---------------------------------------------------------------------------------------------------------------------------------
### Merge
#---------------------------------------------------------------------------------------------------------------------------------

#with tab3:
#if page == "merge":
elif st.session_state.page == 'pdf_mer':    

            st.markdown("---")    
            col_home, title= st.columns([2,15,])
            
            with col_home:
                if st.button("Home", icon="🏠", key="home_fd", type="secondary", use_container_width=True):
                    go_home()
                    st.rerun()

            with title:
                st.info("""The **Extract** tab is designed to extract text and metadata from PDF files. You can upload a PDF, and the tool will pull out the text content for further analysis or editing.""")

            col1, col2 = st.columns((0.2,0.8))
            with col1:           
                with st.container(border=True):              
            
                    uploaded_files = st.file_uploader("**:blue[Choose PDF files]**", type="pdf", accept_multiple_files=True)
                    if uploaded_files:
                        st.success(f"You have selected **{len(uploaded_files)} PDF file(s)** for merging.")
                        if st.button("**Merge PDFs**"):
                            with st.spinner("Merging PDFs..."):
                                merged_pdf = merge_pdfs(uploaded_files)
                            st.success("PDFs merged successfully!")

                            with col2:
                                #st.subheader("**View : Merged PDF**",divider='blue')    
                                with st.container(height=650,border=True):

                                    merged_pdf.seek(0)
                                    images = pdf_to_images(merged_pdf)
                                    for page_num, img in enumerate(images):
                                        st.image(img, caption=f"Page {page_num + 1}", use_container_width=True)
                                
                                st.download_button(label="**📥 Download Merged PDF**",data=merged_pdf,file_name="merged_pdf.pdf",mime="application/pdf")
                    
                    else:
                        st.warning("Please upload PDF files to Merge.")

#---------------------------------------------------------------------------------------------------------------------------------
### Compress
#---------------------------------------------------------------------------------------------------------------------------------

#with tab4:
#if page == "compress":
elif st.session_state.page == 'pdf_comp':

    st.markdown("---")    
    col_home, title= st.columns([2,15,])
            
    with col_home:
        if st.button("Home", icon="🏠", key="home_fd", type="secondary", use_container_width=True):
            go_home()
            st.rerun()

    with title:
        st.info("""The **Compress/Resize** tab is used to reduce the file size of PDF documents. This is useful when you need to share files with size restrictions or save storage space.""") 
        
    col1, col2 = st.columns((0.2,0.8))
    with col1:           
            with st.container(border=True): 
                
                uploaded_file = st.file_uploader("**:blue[Choose PDF file]**", type="pdf",key="file_uploader_resize")
                if uploaded_file is not None:
                
                    st.success(f"You have selected **{uploaded_file.name}** for compress/resize. Please choose the scale factor and press **Compress/Resize**.")
                    st.divider()
                    scale_factor = st.slider("**:blue[Select scale factor]**", 0.1, 2.0, 0.8, 0.1)

                    if st.button("**Compress/Resize**"): 
                        with col2:
                            with st.container(border=True):       
                                
                                try:
                                    images = pdf_to_images(uploaded_file)
                                    resized_images = [img.resize((int(img.width * scale_factor), int(img.height * scale_factor))) for img in images]
                                    resized_pdf = io.BytesIO()
                                    resized_images[0].save(resized_pdf, save_all=True, append_images=resized_images[1:], format="PDF")
                                    resized_pdf.seek(0)

                                    st.success("PDF resized or rescaled successfully!")
                        
                                except Exception as e:
                                    st.error(f"An error occurred: {e}")
                            
                            st.download_button(label="**📥 Download Resized PDF**", data=resized_pdf, file_name="resized_pdf.pdf", mime="application/pdf")
                else:
                    st.warning("Please upload a PDF file to compress or resize.")

#---------------------------------------------------------------------------------------------------------------------------------
### Protect
#---------------------------------------------------------------------------------------------------------------------------------

#with tab5:
#if page == "protect":
elif st.session_state.page == 'pdf_pro':

        st.markdown("---")            
        col_home, title= st.columns([2,15,])
            
        with col_home:
            if st.button("Home", icon="🏠", key="home_fd", type="secondary", use_container_width=True):
                go_home()
                st.rerun()

        with title:  
            st.info("""The **Protect** tab enables you to add password protection to your PDF files. You can set a password to prevent unauthorized access or editing of your documents.""")
        
        col1, col2 = st.columns((0.2,0.8))
        with col1:           
            with st.container(border=True):
                
                uploaded_file = st.file_uploader("**:blue[Choose PDF file]**", type="pdf",key="file_uploader_protect")
                if uploaded_file is not None:

                    st.success(f"You have selected **{uploaded_file.name}** to protect. Please enter the password below and press **Protect** to protect the PDF.")
                    st.divider()
                    password = st.text_input("**:blue[Enter a password to protect your PDF]**", type="password")

                    if st.button("**Protect**"):
                        if password:
                            
                            with col2:
                                with st.container(border=True):
                                    
                                    pdf_reader = PdfReader(uploaded_file)
                                    pdf_writer = PdfWriter()
                                    for page in pdf_reader.pages:
                                        pdf_writer.add_page(page)
                                    with st.spinner("Protecting PDFs..."):
                                        pdf_writer.encrypt(user_pwd=password, owner_pwd=None, use_128bit=True)

                                    output_pdf_io = io.BytesIO()
                                    pdf_writer.write(output_pdf_io)
                                    output_pdf_io.seek(0)  # Move to the start of the file
                                    
                                    st.success(f"Your PDF has been password protected and is ready for download.")
                                st.download_button(label="**📥 Download Password Protected PDF**",data=output_pdf_io,file_name="protected.pdf",mime="application/pdf")

                        else:
                            st.warning("Please enter a password to protect your PDF.")
                else:
                    st.warning("Please upload a PDF file to protect.")

#---------------------------------------------------------------------------------------------------------------------------------
### Unlock
#---------------------------------------------------------------------------------------------------------------------------------

#with tab6:
#if page == "unlock":    
elif st.session_state.page == 'pdf_un':

        st.markdown("---")
        col_home, title= st.columns([2,15,])
            
        with col_home:
            if st.button("Home", icon="🏠", key="home_fd", type="secondary", use_container_width=True):
                go_home()
                st.rerun()

        with title:   
            st.info("""The **Unlock** tab allows you to remove password protection from PDF files. If you have a secured PDF and you know the password, you can unlock it for easier access.""")

        col1, col2 = st.columns((0.2,0.8))
        with col1:           
            with st.container(border=True):     
        
                uploaded_file = st.file_uploader("**:blue[Choose PDF file]**", type="pdf",key="file_uploader_unlock")
                if uploaded_file is not None:

                    st.success(f"You have selected **{uploaded_file.name}** for unlock. Please enter the password below and press **Unlock** to remove the password.")
                    st.divider()
                    password = st.text_input("**:blue[Enter the password to unlock the PDF]**", type="password")

                    if st.button("**Unlock**"):
                        if uploaded_file and password:                           
                        
                            with col2:
                                with st.container(border=True):
                                
                                    try:
                                        with pikepdf.open(uploaded_file, password=password) as pdf:
                                            with st.spinner("Unlocking PDFs..."):
                                                output_pdf_io = io.BytesIO()
                                                pdf.save(output_pdf_io)
                                                output_pdf_io.seek(0)

                                            st.success(f"Password has been removed from the PDF and is ready for download.")

                                    except pikepdf.PasswordError:
                                        st.error("Incorrect password. Please try again.")
                                    except Exception as e:
                                        st.error(f"An error occurred: {str(e)}")
                                        
                                st.download_button(label="**📥 Download Unlocked PDF**",data=output_pdf_io,file_name="unlocked.pdf",mime="application/pdf")
                        else:
                            st.warning("Please enter a password to unlock your PDF.")
                else:
                    st.warning("Please upload a protected PDF file to unlock")

#---------------------------------------------------------------------------------------------------------------------------------
### Rotate
#---------------------------------------------------------------------------------------------------------------------------------

#with tab7:
#if page == "rotate":
elif st.session_state.page == 'pdf_rot':

        st.markdown("---")
        col_home, title= st.columns([2,15,])
            
        with col_home:
            if st.button("Home", icon="🏠", key="home_fd", type="secondary", use_container_width=True):
                go_home()
                st.rerun()

        with title:
            st.info("""The **Rotate** tab lets you change the orientation of pages within a PDF file. You can rotate individual pages or the entire document to correct their orientation.""")

        col1, col2 = st.columns((0.2,0.8))
        with col1:           
            with st.container(border=True): 
                
                uploaded_file = st.file_uploader("**:blue[Choose PDF file]**", type="pdf",key="file_uploader_rotate")
                if uploaded_file is not None:
                
                    st.success(f"You have selected **{uploaded_file.name}** for rotate. Please choose the rotation angle and press **Rotate** to make rotation.")
                    st.divider()
                    rotation_angle = st.slider("**:blue[Select rotation angle]**", 0, 360, 90, 90)

                    if st.button("**Rotate**"): 
                        with col2:
                            with st.container(border=True):       

                                reader = PdfReader(uploaded_file)
                                writer = PdfWriter()
                                for page in reader.pages:
                                    page.rotate(rotation_angle)
                                writer.add_page(page)
                                rotated_pdf = io.BytesIO()
                                writer.write(rotated_pdf)
                                rotated_pdf.seek(0)
        
                                st.success("PDF rotated successfully!")
                            st.download_button(label="**📥 Download Rotated PDF**", data=rotated_pdf, file_name="rotated_pdf.pdf", mime="application/pdf")

                else:
                    st.warning("Please upload a PDF file to rotate.")

#---------------------------------------------------------------------------------------------------------------------------------
### Resize
#---------------------------------------------------------------------------------------------------------------------------------

#with tab8:
#if page == "resize":

#---------------------------------------------------------------------------------------------------------------------------------
### Convert
#---------------------------------------------------------------------------------------------------------------------------------

#with tab9:
#if page == "convert":
elif st.session_state.page == 'pdf_con':

        st.markdown("---")
        col_home, title= st.columns([2,15,])
            
        with col_home:
            if st.button("Home", icon="🏠", key="home_fd", type="secondary", use_container_width=True):
                go_home()
                st.rerun()

        with title:
     
            st.info("""The **Convert** tab offers conversion options between PDF and other formats, such as Word or images. You can upload a PDF and convert it to a different format or vice versa.""")

        col1, col2 = st.columns((0.2,0.8))
        with col1:           
            with st.container(border=True): 
        
                uploaded_file = st.file_uploader("**:blue[Choose PDF file]**", type="pdf", key="file_uploader_convert")
                if uploaded_file is not None:

                    st.success(f"You have selected **{uploaded_file.name}** for conversion.")
                    st.divider()
                    conversion_type = st.selectbox("**:blue[Choose the output format]**", ("Word Document (.docx)", "Plain Text (.txt)"))
                    
                    if st.button("**Convert**"):
                        with col2:
                            with st.container(border=True):       

                                output_file_name = os.path.splitext(uploaded_file.name)[0]
                                
                                if conversion_type == "Word Document (.docx)":
                                    word_file_io = io.BytesIO()  # Use BytesIO instead of writing to disk
                                    doc = Document()
                                    with pdfplumber.open(uploaded_file) as pdf:
                                        with st.spinner("Converting to Word..."):
                                            for page in pdf.pages:
                                                text = page.extract_text()
                                                if text:
                                                    doc.add_paragraph(text)
                                                doc.add_paragraph("\n--- Page Break ---\n")

                                    doc.save(word_file_io)
                                    word_file_io.seek(0)

                                    st.success("PDF converted to Word successfully!")
                                    st.download_button(label="**📥 Download Word File**", data=word_file_io, file_name=f"{output_file_name}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")                               
                                
                                elif conversion_type == "Plain Text (.txt)":
                                    text_file_io = io.BytesIO()                               
                                
                                    with pdfplumber.open(uploaded_file) as pdf:
                                        with st.spinner("Converting to Text..."):
                                            extracted_text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
                                            text_file_io.write(extracted_text.encode("utf-8"))
                                            text_file_io.seek(0)                               
                                
                                    st.success("PDF converted to Text successfully!")
                                    st.download_button(label="**📥 Download Text File**", data=text_file_io, file_name=f"{output_file_name}.txt", mime="text/plain")

                else:
                    st.warning("Please upload a PDF file to convert.")

#---------------------------------------------------------------------------------------------------------------------------------
### Summarization
#---------------------------------------------------------------------------------------------------------------------------------

#if page == "summary":
elif st.session_state.page == 'pdf_sum':

        st.markdown("---")
        col_home, title= st.columns([2,15,])
            
        with col_home:
            if st.button("Home", icon="🏠", key="home_fd", type="secondary", use_container_width=True):
                go_home()
                st.rerun()

        with title:   
            st.info("""The **Summarization** tab offers summary of the uploaded pdf without using GenAI.""")

        col1, col2 = st.columns((0.2,0.8))
        with col1:           
            with st.container(border=True): 
        
                uploaded_file = st.file_uploader("**:blue[Choose PDF file]**", type="pdf", key="file_uploader_convert")
                if uploaded_file is not None:
                        
                    with pdfplumber.open(uploaded_file) as pdf:
                        total_pages = len(pdf.pages)
                    st.success(f"📂 File uploaded: **{uploaded_file.name}**")
                    st.divider()
                    
                    #num_sentences = st.slider(":blue[Select number of sentences for summary]", 1, 10, 5)
                    summary_choice = st.radio("**:blue[Choose to summarize]**", ["All Pages", "Specific Page", "Range of Pages"])

                    start_page, end_page = None, None
                    if summary_choice == "Specific Page":
                        start_page = st.number_input("**:blue[Enter page number]**", min_value=1, max_value=total_pages, value=1)
                        end_page = start_page  # Only one page selected
                    elif summary_choice == "Range of Pages":
                        start_page = st.number_input("**:blue[Start Page]**", min_value=1, max_value=total_pages, value=1)
                        end_page = st.number_input("**:blue[End Page]**", min_value=start_page, max_value=total_pages, value=total_pages)

                    num_sentences = 10
                    if st.button("**Summarization**"):
                        with col2:
                            with st.container(border=True):  
                                
                                with st.spinner("Extracting text..."):
                                    with st.spinner("Summarizing..."):
                                        extracted_text = extract_text_from_pdf(uploaded_file, start_page, end_page)
                                        summary = summarize_text(extracted_text, num_sentences)
                                        
                                        formatted_summary = f"📄 **Summary of {uploaded_file.name}**\n\n{summary}"
                                        st.text_area("", formatted_summary, height=200)
                                        summary_file = io.BytesIO(summary.encode("utf-8"))
                                        
                            st.download_button(label="**📥 Download Summary (.txt)**",data=summary_file,file_name="summary.txt",mime="text/plain")
                                
                    #else:
                        #st.warning("No text could be extracted from the PDF.")

                else:
                    st.warning("Please upload a PDF file to summarize.")                                
                 
#---------------------------------------------------------------------------------------------------------------------------------


